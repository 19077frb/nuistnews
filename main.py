from requests_html import HTML, HTMLSession, UserAgent
from PIL import Image
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches
import os

# proxy = {
#     'http': '127.0.0.1:62201'
# }
# os.environ['NO_PROXY']='news.nuist.edu.cn'
def open_url(url):
    # random_agent = USER_AGENTS[randint(0, len(USER_AGENTS) - 1)]
    random_agent = UserAgent().random
    header = {
        'User-Agent': random_agent,
        'Connection': 'close',
    }
    session = HTMLSession()
    resp = session.get(url, headers=header)
    resp.html.render(sleep=60, wait=60, timeout=300)
    # resp.html.render()
    return resp.html.html


folder = os.getcwd() + '\\pic\\'


if not os.path.exists(folder):
    os.makedirs(folder)

folder = os.getcwd() + '\\result\\'


if not os.path.exists(folder):
    os.makedirs(folder)

url = str(input("请输入网址\n"))
# url = 'https://news.nuist.edu.cn/_web/search/doSearch.do?locale=zh_CN&request_locale=zh_CN&_p=YXM9MTcmdD03NyZkPTIzMSZwPTMmZj0xNzQmYT0wJm09U04mfGJubkNvbHVtblZpcnR1YWxOYW1lPTE3NCY_'
# url = 'http://news.nuist.edu.cn/2023/0824/c1147a227291/page.htm'
print("start")
response = open_url(url)

# print(response)
print("success get html")

# fh = open("./urllib_test_runoob_search.html", "w", encoding="utf-8")  # 将文件写入到当前目录中
# # fh = open("./urllib_test_runoob_search.html", "w")  # 将文件写入到当前目录中
# fh.write(response)
# fh.close()

html = HTML(html=response)
print(html.absolute_links)
num = 0
for link in html.absolute_links:
    if 'htm' not in link:
        continue
    num += 1
    document = Document()

    news_html = open_url(link)
    news_html = HTML(html=news_html)

    news_title = news_html.search('<h1 class="arti-title">{}</h1>')[0]

    document.add_heading(news_title, level=0)
    print(news_title)

    article = news_html.find('div.wp_articlecontent')[0]

    paragraphs = article.find('p')
    cnt = 0
    for paragraph in paragraphs:
        if 'src' in paragraph.html:
            cnt += 1
            # print(paragraph.search('src="{}"'))
            img_url = 'http://news.nuist.edu.cn' + paragraph.search('src="{}"')[0]
            req = requests.get(img_url)
            image = Image.open(BytesIO(req.content))
            # fileName = str(uuid.uuid4()) + '.' + image.format.lower()
            fileName = news_title + '_' + str(cnt)
            with open('./pic/' + fileName, 'wb') as f:
                f.write(req.content)
            document.add_picture('./pic/' + fileName, width=Inches(6))
        else:
            # print(paragraph.text)
            document.add_paragraph(paragraph.text)

    document.save('./result/' + news_title + '.docx')
    print("结束,共获取" + str(num) + "篇文章，请到result文件夹中查看")
